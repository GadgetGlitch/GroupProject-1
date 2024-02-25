#!/usr/bin/env python
# coding: utf-8

# # 1. Install necessary packages and libraries
# Run the following four installs only once

# In[ ]:


#!pip install geopy
#!pip install tzwhere
#!pip install timezonefinder
#!pip install pytz


# In[1]:


# 1. Setup and Imports
import pandas as pd
import pytz
import plotly.graph_objects as go
import pandas as pd
import ipywidgets as widgets
import datetime
import ipywidgets as widgets
import fitz  # PyMuPDF for PDF files
import pandas as pd
import requests
import re  # For regular expressions
import openpyxl  # For .xlsx files
import xlrd  # For .xls files
from docx import Document
from bs4 import BeautifulSoup
from IPython.display import display
from geopy.geocoders import Nominatim
from timezonefinder import TimezoneFinder


# In[4]:


# 1. See install and import section

# Read CSV and convert DATE column to PREDeadline Dates and Timezones
# Load Original CSV Dataset file
file_path = r"C:\Users\annak\CAS-502 Computation\GroupProject-1\GP1 Dataset_ Assignments.csv"
df = pd.read_csv(file_path, encoding='ISO-8859-1')


# Initialize Nominatim and TimezoneFinder
geolocator = Nominatim(user_agent="geocity")
tf = TimezoneFinder()

def get_timezone_from_city(city):
    location = geolocator.geocode(city)
    if location:
        timezone_str = tf.timezone_at(lat=location.latitude, lng=location.longitude)
        return timezone_str
    return None

# Function to standardize city names
def standardize_city_name(city):
    city = city.lower()
    if city in ["saranda", "sarande"]:
        return "Sarandë, Albania"
    return city

def get_predeadline(deadline, buffer_days=1):
    """Subtract buffer days from deadline date to get the pre-deadline date."""
    return deadline - datetime.timedelta(days=buffer_days)

# Read Dates and Establish Date Format
df['DATE'] = pd.to_datetime(df['DATE'], format='%m/%d/%Y')

# Function to calculate pre-deadline date
def get_predeadline(deadline, buffer_days=1):
    predeadline = deadline - datetime.timedelta(days=buffer_days)
    return predeadline

# Calculate predeadline for each date
df['PREDEADLINE'] = df['DATE'].apply(lambda x: get_predeadline(x, buffer_days=1))



# Rearrange columns to put PREDEADLINE between DATE and WK#
column_order = ['ITEM#', 'DATE', 'DEADLINE', 'TIMEZONE', 'PREDEADLINE', 'DUETIME', 'WK#', 'COURSE', 'ASSIGNMENT NAME', 'DESCRIPTION', 'DONE?']
df = df[column_order]

# Print the new table headers
print(df.head())


# In[ ]:


#Convert due dates to the user's local timezone and calculate for pre-deadline dates
# See import install section above

geolocator = Nominatim(user_agent="geocity")
tf = TimezoneFinder()

#Asing Student for their city for the timezone change if necessary
user_city = standardize_city_name(input("Enter the city you are in: "))
user_timezone = get_timezone_from_city(user_city)


# Function to get timezone from city name
def get_timezone_from_city(city):
    """Return the timezone of the city."""
    location = geolocator.geocode(city)
    if location:
        timezone_str = tf.timezone_at(lat=location.latitude, lng=location.longitude)
        return timezone_str
    return None

# Function to standardize city names (e.g., handle variations)
def standardize_city_name(city):
    city = city.lower()
    if city in ["saranda", "sarande"]:
        return "Sarandë, Albania"
    return city

# Load CSV and convert DATE column to datetimes
df = pd.read_csv("GP1 Dataset_ Assignments.csv", encoding='ISO-8859-1')
df['DATE'] = pd.to_datetime(df['DATE'], utc=True)

# Ask the user for their city
user_city = standardize_city_name(input("Enter the city you are in: "))
user_timezone = get_timezone_from_city(user_city)

# Convert assignment dates to the user's local timezone
if user_timezone:
    user_tz = pytz.timezone(user_timezone)
    df['LOCAL_DUE_DATE'] = df['DATE'].apply(lambda x: x.astimezone(user_tz))
else:
    print("City not found. Using UTC as default.")
    df['LOCAL_DUE_DATE'] = df['DATE']

    
print(df[['DATE', 'LOCAL_DUE_DATE']])



# In[ ]:


# Rearrange columns to include calculated Timze zone and P
column_order = ['ITEM#', 'DATE', 'DEADLINE', 'TIMEZONE', 'PREDEADLINE', 'DUETIME', 'WK#', 'COURSE', 'ASSIGNMENT NAME', 'DESCRIPTION', 'DONE?']
df = df[column_order]


# In[ ]:


# Make a new table file that can output the new synthetic data and columns. 

new_file_path = file_path.rsplit('\\', 1)[0] + '\\GP1 Dataset_ Assignments_PreDeadline.csv'
df.to_csv(new_file_path, index=False)
print(f"File saved as: {new_file_path}")


# # Failed Nice-toHave: Attempt to to combine functions and read files once uploaded from doc/pdf/excel/url

# In[ ]:


# Failed Nice-toHave: Attempt to to combine functions and read files once uploaded from doc/pdf/excel/url¶
# Working UPLOAD BUTTON

# 1. See Install import section
# 2. Utility Functions

def extract_info_from_text(text):
    due_date_match = re.search(r"Due Date: (\d{2}/\d{2}/\d{4})", text)
    due_date = due_date_match.group(1) if due_date_match else "Not found"
    # Implement similar searches for class number/code, assignment details, etc.
    return {'due_date': due_date}

def process_docx(file_content):
    doc = Document(file_content)
    text = "\n".join([para.text for para in doc.paragraphs])
    return extract_info_from_text(text)

def process_pdf(file_content):
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ''.join([page.get_text("text") for page in doc])
        return extract_info_from_text(text)
    except Exception as e:
        print(f"Failed to process PDF file: {e}")
        return {}

def process_excel(file_content, file_type='xlsx'):
    df = pd.read_excel(file_content, engine='openpyxl' if file_type == 'xlsx' else 'xlrd')
    return df.to_dict('records')

def process_csv(file_content):
    df = pd.read_csv(file_content)
    return df.to_dict('records')

def scrape_and_process_html(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        text = soup.get_text()
        return extract_info_from_text(text)
    except Exception as e:
        print(f"Failed to scrape HTML from {url}: {e}")
        return {}

def process_file_upload(change):
    uploaded_files = file_upload.value  # Correctly access the uploaded files
    for name, file_info in uploaded_files.items():
        content = file_info['content']
        if name.endswith('.xlsx') or name.endswith('.xls'):
            # Process Excel files
            assignments = process_excel(content, 'xlsx' if name.endswith('.xlsx') else 'xls')
        elif name.endswith('.csv'):
            # Process CSV files
            assignments = process_csv(content)
        elif name.endswith('.pdf'):
            # Process PDF files
            assignments = process_pdf(content)
        elif name.endswith('.docx'):
            # Process DOCX files
            assignments = process_docx(content)
        else:
            print(f"Unsupported file type: {name}")
            assignments = []
       
        
# 3. Set up the widgets for file upload and URL input
file_upload = widgets.FileUpload(accept='.xlsx, .xls, .pdf, .docx', multiple=True, description='Upload Files')
file_upload.observe(process_file_upload, names='value')

url_input = widgets.Text(placeholder='Enter URL here', description='URL:')
submit_button = widgets.Button(description='Submit URL', button_style='', tooltip='Click to submit the URL')

def submit_url(b):
    url = url_input.value
    if url:
        result = scrape_and_process_html(url)
        print(result)  # Display or handle the scraped data
    else:
        print("Please enter a URL.")

submit_button.on_click(submit_url)

# 4. Organize and display UI components
upload_box = widgets.VBox([widgets.Label('Supported file types: Docs (Word), PDF, Excel'), file_upload])
url_box = widgets.VBox([url_input, submit_button])
main_box = widgets.VBox([upload_box, url_box])
display(main_box)


# In[ ]:


# Failed: ISSUE 20 Try to read PDF/DOC/EXCEL/URL
# Processing functions for file upload
def process_file_upload(change):
    uploaded_files = change['new']
    for name, file_info in uploaded_files.items():
        # Detect file type and process accordingly
        if name.endswith('.xlsx') or name.endswith('.xls'):
            df = pd.read_excel(file_info['content'])
            print(df.head())  # Example of processing Excel file
        elif name.endswith('.pdf'):
            doc = fitz.open(stream=file_info['content'], filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            print(text)  # Example of processing PDF file
        elif name.endswith('.docx'):
            doc = Document(file_info['content'])
            text = ""
            for para in doc.paragraphs:
                text += para.text + '\n'
            print(text)  # Example of processing Word document
        else:
            print(f"Unsupported file type: {name}")

# Function to scrape and process HTML from URL
def scrape_and_process_html(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    print(soup.get_text())  # Example of processing HTML content

# Create the file upload widget
file_upload = widgets.FileUpload(
    accept='.xlsx, .xls, .pdf, .docx',  # Specify acceptable file types
    multiple=True,  # Allow multiple files to be selected
    description='Upload Files'
)
file_upload.observe(process_file_upload, names='value')

# Text label with supported file types
file_types_label = widgets.Label('Supported file types: Docs (Word), PDF, Excel')

# URL input and submission setup
url_input = widgets.Text(placeholder='Enter URL here', description='URL:')
submit_button = widgets.Button(description='Submit URL', button_style='', tooltip='Click to submit the URL')
def submit_url(b):
    url = url_input.value
    if url:  # Check if the URL is not empty
        scrape_and_process_html(url)
    else:
        print("Please enter a URL.")
submit_button.on_click(submit_url)

# Organizing the layout
upload_box = widgets.VBox([file_types_label, file_upload])
url_box = widgets.VBox([url_input, submit_button])
main_box = widgets.VBox([upload_box, url_box])

# Displaying the entire layout
display(main_box)



# # MAIN APP CODE

# In[ ]:


#1 See import install section

#2 Read csv file dataset
df = pd.read_csv('GP1 Dataset_ Assignments.csv')

#3 Action before and after item is complete. Text starts Black when complete Red.
df['Colors'] = 'nil'
for index,row in df.iterrows():
    if row['Status'] == 'Open':
        df.at[index,'Colors'] = 'black'
    else:
        df.at[index,'Colors'] = 'red'

#4 What data to scrape from CSV
a_list = list(df['ASSIGNMENT NAME'])
a_list.insert(0,'None')

#5 Dropdown button that allows student to mark an assignment compelte
button = widgets.Button(description="Mark As Complete")
text = widgets.Dropdown(
    options=a_list,
    value=a_list[0],
    description='Assignment Name:',
    disabled=False,
)

output = widgets.Output()

#6 Triggers changes to the UI when done Y from N
#@output.capture()
def on_button_clicked(b):
    with output:
        output.clear_output()
        a_name = text.value
        
        if a_name != 'None':
            
       
            index_req = df.index[df['ASSIGNMENT NAME'] == a_name].to_list()
        
            df.at[index_req[0],'Status'] = 'Y'
            df.at[index_req[0],'Colors'] = 'red'
            df.sort_values(by=['Status', 'DATE'], inplace=True, ascending=[False,True])
            a_list = list(df['ASSIGNMENT NAME'])
        
        df.to_csv('GP1 Dataset_ Assignments.csv',columns=['DATE','PREDEADLINE','TIME', 'WEEK NUMBER','COURSE','ASSIGNMENT NAME','DESCRIPTION','Status'])
        fig = go.Figure(data=[go.Table(header=dict(
        values=["Date", "Time","PreDeadline", "Week Number","Course","Assignment Name","Description","Status"],
        line_color='white', fill_color='white',
        align='center', font=dict(color='black', size=12)
        ),
        cells=dict(
        values=[df['DATE'], df['WEEK NUMBER'],df['COURSE'],df['ASSIGNMENT NAME'],df['DESCRIPTION'],df['Status']],
        #line_color=[df.Colors], 
        #fill_color=[df.Colors],
        align='center', font=dict(color= [df.Colors], size=11)
         ))
        ])
        fig.show()

        
fig = go.Figure(data=[go.Table(header=dict(
        values=["Date", "Week Number","Course","Assignment Name","Description","Status"],
        line_color='white', fill_color='white',
        align='center', font=dict(color='black', size=12)
          ),
        cells=dict(
        values=[df['DATE'], df['WEEK NUMBER'],df['COURSE'],df['ASSIGNMENT NAME'],df['DESCRIPTION'],df['Status']],
        #line_color=[df.Colors], 
        #fill_color=[df.Colors],
        align='center', font=dict(color= [df.Colors], size=11)
          ))
        ])


#7 Widgets and UI
button.on_click(on_button_clicked)
list_button = widgets.HBox([text,button])
widgets.VBox([list_button,output])



